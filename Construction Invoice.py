import glob
import os
import pandas as pd
import datetime as dt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import docx
from docx.shared import Inches
from docxtpl import DocxTemplate

path = os.path.expanduser('/Users/xintongli/PycharmProjects/Automatic Invoice Project/Dataset/*.xlsx')
#path1 ='/Users/xintongli/Documents/学习/Python/Interested Project/Construction Invoice/*.xlsx'
data = {}
#pd.set_option('precision', 2) this does not change the export to docx
#Loop over the folder to find the file that has the xlsx suffix
for i,f in enumerate(glob.glob(path)):
    data[f[88:(len(f)-5)]] = pd.read_excel(f)
'''Also work, no need to add the enumerate
data1 = {}
for f in glob.glob(path):
    data1[f[88:(len(f)-5)]] = pd.read_excel(f)'''

today = dt.date.today()
def projectdata():
    '''Create a dictionary that contains data from individual firm-key:project name, value:project detail'''
    for key in data.keys():
        data[key].date = pd.to_datetime(data[key].date, format='%m/%d')
        data[key].date = data[key].date.apply(lambda dt: dt.replace(year=today.year))
    #Create working hour
        data[key]['wtime'] = data[key].time_end - data[key].time_start
        #data[key]['hwage'] = np.random.uniform(40,50,len(data[key].index))
        data[key]['dinditotal'] = data[key].wtime * data[key].hwage
        data[key]['construction'] = str(key)
        data[key] = data[key].round(2)
        wdata = {}
        wdata[key] = pd.pivot_table(data[key], values='dinditotal', index='name', columns='date')
        wdata[key] = wdata[key].fillna(0)
        wdata[key]['winditotal'] = wdata[key].sum(axis=1)
        wdata[key]['totalpay'] = wdata[key]['winditotal'].cumsum(axis=0)
    return data, wdata
data, wdata = projectdata()


def webdata(data):
    '''create job_web based data dictionary, in order to export as docx'''
    tdata = pd.DataFrame()
    for value in data.values():
        tdata = pd.concat([tdata, value],ignore_index=True)
    tdata = tdata.groupby(['name','job_web']).agg({'dinditotal': sum,
                                                'wtime': sum,
                                                'hwage': 'mean'})
    t = tdata.groupby(['job_web','name']).agg({'dinditotal': sum,
                                                'wtime': sum,
                                                'hwage': 'mean'})
    t['totalpay'] = t.groupby(level=0)['dinditotal'].cumsum()
    namelist = list(set(t.index.get_level_values(0)))
    df = {}
    for job_web in namelist:
        df[job_web] = t.loc['{}'.format(job_web)]
    return df
df = webdata(data)
#Create Docx
header = ['Name','Wage','Hourly Wage','Hour']
#Use docx to export as table in the word
'''def docxway(key):
    for key in df.keys():
        doc = docx.Document()
        head = doc.add_heading('INVOICE')
        parabill = doc.add_paragraph('Bill To: '+'{}'.format(key))
        r.add_picture('logo.png', width=Inches(3), height=Inches(1.5))
        paraivnum = doc.add_paragraph('Invoice #: '
                                   +'{}'.format(key)
                                   +'{}'.format(today.strftime('%Y%m%d')))
        paradate = doc.add_paragraph('Invoice Date: '+'{}'.format(today))
        paradue = doc.add_paragraph('Due Date: '+'{}'.format(today+dt.timedelta(days=7)))
        table = doc.add_table(rows=len(df[key].index)+1, cols=(len(df[key].columns)))
        for j in range(len(header)):
            table.cell(0,j).text = header[j]
        for i in range(df[key].shape[0]):
            table.cell(i+1,0).text = str(df[key].index[i])
        for i in range(df[key].shape[0]):
            for j in range(len(header)-1):
                table.cell(i+1,j+1).text = str(df[key].iloc[i,j])
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        paratotalpay = doc.add_paragraph('Total Pay: '+'{}'.format(df[key].iloc[-1,-1]))
        paratotalpay.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.save('/Users/xintongli/Documents/学习/Python/Interested Project/Construction Invoice/Invoice/'+'{}'.format(key)+'_invoice.docx')
'''
#Use the docxtemplate to export as doc--no need to add picture
def docfile(key):
    tpl = DocxTemplate('/Users/xintongli/Documents/学习/Python/Interested Project/Construction Invoice/Invoice_template.docx')
    context = {'parabill': '{}'.format(key),
               'paraivnum': '{}'.format(key) + '{}'.format(today.strftime('%Y%m%d')),
               'paradate': '{}'.format(today),
               'paradue':'{}'.format(today+dt.timedelta(days=7)),
               'paratotalpay':'{}'.format(df[key].iloc[-1,-1])}
    table = tpl.add_table(rows=len(df[key].index) + 1, cols=(len(df[key].columns)))
    for j in range(len(header)):
        table.cell(0, j).text = header[j]
    for i in range(df[key].shape[0]):
        table.cell(i + 1, 0).text = str(df[key].index[i])
    for i in range(df[key].shape[0]):
        for j in range(len(header) - 1):
            table.cell(i + 1, j + 1).text = str(df[key].iloc[i, j])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    paratotalpay = tpl.add_paragraph('Total Pay: CA$'+'{}'.format(df[key].iloc[-1,-1]))
    paratotalpay.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tpl.render(context)
    tpl.save('/Users/xintongli/PycharmProjects/Automatic Invoice Project/Dataset/Invoice/'
             +'{}'.format(key)+'_invoice.docx')

for key in df.keys():
    docfile(key)

